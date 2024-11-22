# python iDocumentosv002.py
import customtkinter as ctk
from tqdm import tqdm
from tkinter import messagebox
from docx import Document
from tkinter import filedialog
from tkinter import messagebox
from datetime import datetime
from docx2pdf import convert
from PyPDF2 import PdfMerger # nao depende
from tkinter import PhotoImage
import pandas as pd
import os

#A conversão para PDF só funciona se o Word estiver instalado na máquina.

class App():
    #lista completa de alunos
    tabAlunos = None
    #lista completa de turmas
    tabTurmas = None
    #tabela de turmas e quantidade de alunos
    tabinfo = None
    #array que contêm turmas
    arrayinfo = None
    #pasta de destino
    pastaDestino = None

    tela = None
    def objeto(self, telas):
        self.tela = telas

    def selecionar_tabela(self):
        try:
            arquivo = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
            if arquivo:
                self.tabAlunos = pd.read_excel(arquivo)
                self.tabTurmas = pd.read_excel(arquivo, sheet_name=1)
            else:
                messagebox.showwarning("Nenhum arquivo Selecionado", "A ação será encerrada.") 
                return "erro"
        except:
            messagebox.showwarning("Erro ao selecionar planilha", "A ação será encerrada.")
            print("Erro ao escolher arquivo")

    # a função abaixo pergunta a pasta de destino
    def pasta_destino(self):
        try:
            self.pastaDestino = filedialog.askdirectory()
            print(self.pastaDestino)
            if not self.pastaDestino:  # Verifica se nenhuma pasta foi selecionada
                messagebox.showwarning("Nenhuma Pasta Selecionada", "A ação será encerrado.")
                return "erro"
            if self.pastaDestino:
                messagebox.showinfo("Pasta Selecionada", f"Os arquivos serão salvos em:\n{self.pastaDestino}")
        except Exception as e:
            print("Erro ao selecionar Pasta de Destino: ", e)
    
    def info(self):
        try:
            tabinfo = self.tabAlunos.groupby('Turma').size().reset_index(name='quantidade_alunos')
            self.tabinfo = tabinfo.sort_values(by='Turma', ascending=False)
            self.arrayinfo = self.tabinfo['Turma'].tolist()
        except:
            print("Erro ao analisar tabela")

    def gerarArquivoPorAluno(self, template, turma, pdf):
        try: # se a turma não for especificada ele gera de todos os estudantes
            tabela = self.tabAlunos
            if turma != None:
                #agrupa os dados dos estudantes que estão na mesma turma em uma nova tabela
                tabela = self.tabAlunos.query(f'Turma == "{turma}"')
                tabela = tabela.sort_values(by='Nome')
            arquivo = Arquivo()
            nomeArquivos = arquivo.arquivoPorAluno(tabela,template,self.pastaDestino)
            print(nomeArquivos)
            if pdf:
                arquivo.converterParaPdf(nomeArquivos, self.pastaDestino, turma)
            print("finalizado")
        except Exception as e:
            print("Erro ao gerar Arquivos:", e)
    
    def gerarArquivoPorTurma(self, template, turma, pdf, mes1, mes2):
        try: # apenas uma turma
            nomeArquivos = []
            tabela = self.tabAlunos
            if turma != None:
                #agrupa os dados dos estudantes que estão na mesma turma em uma nova tabela
                tabela = self.tabAlunos.query(f'Turma == "{turma}"')
                tabela = tabela.sort_values(by='Nome')
                arquivo = Arquivo()
                nomeArquivos = arquivo.arquivosPorTurma(tabela,template,self.pastaDestino, mes1, mes2)
                print(nomeArquivos)
                print("Arquivos gerados com sucesso")
                if pdf:
                    arquivo.converterParaPdf(nomeArquivos, self.pastaDestino, turma)
                print("finalizado")
            else: # todas as turma
                for linha in self.tabTurmas.index:
                    turma = str(self.tabTurmas.loc[linha, "Turma"])
                    tabela = None
                    tabela = self.tabAlunos.query(f'Turma == "{turma}"')
                    tabela = tabela.sort_values(by='Nome').reset_index(drop=True)
                    if tabela.empty:
                        continue
                    print(turma)
                    print(tabela)
                    arquivo = Arquivo()
                    nomeArquivo = arquivo.arquivosPorTurma(tabela, template, self.pastaDestino, mes1, mes2)
                    nomeArquivos.append(nomeArquivo)
                print(nomeArquivos)
                print("Arquivos gerados com sucesso")
                if pdf:
                    arquivo.converterParaPdf(nomeArquivos, self.pastaDestino, nomeArquivos[0])
                print("finalizado")
        except Exception as e:
            print("Erro ao gerar Arquivos:", e)

    def gerarArquivoIrmaos(self, template, pdf):
        try: 
            alunos = self.tabAlunos
            arquivo = Arquivo()
            tabAlunos =arquivo.procurarIrmaos(alunos)
            nomeArquivo = arquivo.arquivosPorTurma(tabAlunos, template, self.pastaDestino, "irmao", None)
            turma = "PDF-irmaos"
            if pdf:
                arquivo.converterParaPdf(nomeArquivo, self.pastaDestino, turma)
        except Exception as e:
            print("Erro ao gerar Arquivos:", e)


class Arquivo:
    #template que vai ser utilizado
    template = None
    #Turma que vai ser gerado
    turma = None
    #tabela de nomes
    tabelaAlunos = None
    #caminho de slavamento
    pastaDestino = None

    def procurarIrmaos(self, alunos):
        try:
            tbirmao = pd.DataFrame()
            fmPossuiIrmo = 0
            while not alunos.empty:
                
                #print(alunos.iloc[0])
                mae = alunos.iloc[0]['Mãe']
                mae = mae.strip().lower()
                #print(mae)
                irmaos = alunos[alunos['Mãe'].str.strip().str.lower() == mae]
                print('Procurando irmãos...')
                if(len(irmaos) > 1):
                    tbirmao = pd.concat([tbirmao, irmaos.reset_index(drop=True)])
                    #print(tbirmao)
                    indices_irmaos = irmaos.index
                    alunos = alunos.drop(indices_irmaos)
                    fmPossuiIrmo += 1

                    print('Transferindo irmãos...')
                else:

                    print('irmãos não encontrados')
                    alunos = alunos.drop(alunos.index[0])
            tbirmao = tbirmao.reset_index(drop=True)
            #print(alunos)
            print('Transferido ' + str(fmPossuiIrmo) + ' familias')
            print(str(fmPossuiIrmo) + ' mães possuem mais de um filho')
            print(tbirmao)
            return tbirmao
        except Exception as e:
            print("erro ao procurar irmãos:", e)


    def arquivoPorAluno(self,tabAlunos, template, pastaDestino ):
        try:
            nomeArquivos = []
            for linha in tabAlunos.index:
                documento = Document(template)
                dicionario = {
                    "<<nome>>": str(tabAlunos.loc[linha, "Nome"]),
                    "<<turma>>": str(tabAlunos.loc[linha, "Turma"]),
                    "<<turno>>": str(tabAlunos.loc[linha, "Turno"]),
                    "<<codigo>>": str(tabAlunos.loc[linha, "Código"]),
                    "<<dia>>": str(datetime.now().day),
                    "<<mes>>": str(datetime.now().month),
                    "<<ano>>": str(datetime.now().year),
                }
                if template != "templates/template_Uniforme.docx":
                    header = documento.sections[0].header
                    #percorre todos os paragrafos do header
                    for paragrafo in header.paragraphs:
                        for codigo in dicionario:
                            if codigo in paragrafo.text:
                                paragrafo.text = paragrafo.text.replace(codigo, dicionario[codigo])


                #percorre todos os paragrafos do documento
                for paragrafo in documento.paragraphs:
                    #precorre o dicionario
                    for codigo in dicionario:
                        #se existir o codigo no paragrafo
                        if codigo in paragrafo.text:
                            paragrafo.text = paragrafo.text.replace(codigo, dicionario[codigo])
                nomeArquivo = f"{dicionario['<<nome>>']}"
                documento.save(os.path.join(pastaDestino, f"{nomeArquivo}.docx"))
                nomeArquivos.append(nomeArquivo)
                #convert(os.path.join(pastaDestino, f"{dicionario['<<nome>>']}.docx"), os.path.join(pastaDestino, f"{dicionario['<<nome>>']}.pdf"))
            return nomeArquivos
        except Exception as e:
            print("Erro ao gerar arquivos por Aluno:", e)
    
    def arquivosPorTurma(self,tabAlunos, template, pastaDestino, mes1, mes2):
        nomeArquivos = []
        try:
            documento = Document(template)
            linha = 0
            dicionario = {
                "<<nome>>": str(tabAlunos.loc[linha, "Nome"]) if tabAlunos.loc[linha, "Nome"] is not None else "",
                "<<mes1>>": mes1 if mes1 is not None else "",
                "<<mes2>>": mes2 if mes2 is not None else "",
                "<<mae>>": str(tabAlunos.loc[linha, "Mãe"]) if tabAlunos.loc[linha, "Mãe"] is not None else "",
                "<<turma>>": str(tabAlunos.loc[linha, "Turma"]) if tabAlunos.loc[linha, "Turma"] is not None else "",
                "<<turno>>": str(tabAlunos.loc[linha, "Turno"]) if tabAlunos.loc[linha, "Turno"] is not None else "",
                "<<codigo>>": str(tabAlunos.loc[linha, "Código"]) if tabAlunos.loc[linha, "Código"] is not None else "",
                "<<dia>>": str(datetime.now().day),
                "<<mes>>": str(datetime.now().month),
                "<<ano>>": str(datetime.now().year),
            }
            header = documento.sections[0].header
            #percorre todos os paragrafos do header
            for paragrafo in header.paragraphs:
                for codigo in dicionario:
                    if codigo in paragrafo.text:
                        paragrafo.text = paragrafo.text.replace(codigo, dicionario[codigo])


            #percorre todos os paragrafos do documento
            for paragrafo in documento.paragraphs:
                #precorre o dicionario
                for codigo in dicionario:
                    #se existir o codigo no paragrafo
                    if codigo in paragrafo.text:
                        paragrafo.text = paragrafo.text.replace(codigo, dicionario[codigo])

            #pega a primeira tabela do documento
            tabelaDoDocumento = documento.tables[0]

            cabecalho = [celula.text for celula in tabelaDoDocumento.rows[0].cells]

            #procura e altera as chaves do texto do cabeçalho da tabela do documento de acordo com o dicionario 
            for celula in tabelaDoDocumento.rows[0].cells:
                for codigo in dicionario:
                    celula.text = celula.text.replace(codigo, dicionario[codigo])

            for index in tabAlunos.index:
                #cria uma nova linha
                linhaNova = tabelaDoDocumento.add_row() 
                for coluna in tabAlunos.columns:
                    if coluna in cabecalho:
                        idx = cabecalho.index(coluna)
                        #preencher a célula correspondente na novas linha
                        linhaNova.cells[idx].text = str(tabAlunos.loc[index, coluna]) + "\n"
            if mes1 == "irmao":
                nomeArquivo = "Irmaos"
            else:
                nomeArquivo = f"{dicionario['<<turma>>']}".replace(" ", "")
            documento.save(os.path.join(pastaDestino, f"{nomeArquivo}.docx"))
            nomeArquivos.append(nomeArquivo)
            return nomeArquivos
        except Exception as e:
            print("erro ao gerar arquivos por turma:", e)

    def converterParaPdf(self, nomeArquivos, pastaDestino, turma):
        try:
            for nomeArquivo in nomeArquivos:
                convert(os.path.join(pastaDestino, f"{nomeArquivo}.docx"), os.path.join(pastaDestino, f"{nomeArquivo}.pdf")) 
            juntador = PdfMerger()
            for pdf in nomeArquivos:
                juntador.append(os.path.join(pastaDestino, f"{pdf}.pdf"))
            if turma == None:
                turma = "Estudantes"
            juntador.write(os.path.join(pastaDestino, f"{turma}.pdf"))
            juntador.close()

            for pdf in nomeArquivos:
                caminho_pdf = os.path.join(pastaDestino, f"{pdf}.pdf")
                if os.path.exists(caminho_pdf):
                    os.remove(caminho_pdf)

        except Exception as e:
            print("erro ao converter arquivos para pdf:", e)

        


 
    #gerador = App()
    #gerador.selecionar_tabela()
    #gerador.pasta_destino()
    #gerador.gerarArquivoPorTurma("templates/template_reuniao.docx", "1º ANO - A")  # coloque None na turma para poder gerar de todas as turmas
    #gerador.gerarArquivoPorAluno("templates/template_Uniforme.docx", "1º ANO - A", False)   #coloque None na turma para poder gerar de todas as turmas o true or false caso queira em pdf
    #gerador.gerarArquivoIrmaos()

# Configurando a aparência do CustomTkinter
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("dark-blue") 

# Classe principal que contém todas as telas
class Tela(ctk.CTk):
    
    def __init__(self):
        super().__init__()
        self.iconbitmap("templates/logo66.ico")
        #logo = ctk.CTkImage(file="templates/logo.png")
        #logo_img = PhotoImage(file="templates/logo.png")
        #self.iconphoto(False, logo_img)
        self.title("IDocumentos")
        self.geometry("610x440")
        self.resizable(False, False)
        self.tela_atual = None
        
        self.frame = None
        self.trocar_tela(self.tela_inicial)

    gerador = None
    def objeto(self, app):
        self.gerador = app

    def trocar_tela(self, nova_tela):
        # Remove o frame atual, se houver
        if self.frame is not None:
            self.frame.pack_forget()

        # Cria o novo frame
        self.frame = ctk.CTkFrame(self, fg_color=None)
        self.frame.pack(pady=20, padx=60, expand=True, fill="both")

        # Chama a função para construir a nova tela
        nova_tela(self.frame)


    def tela_inicial(self, frame):
        label_titulo = ctk.CTkLabel(frame, text="IDocumentos", font=ctk.CTkFont(size=20, weight="bold"))
        label_titulo.pack(pady=12, padx=10)

        label_instrucoes = ctk.CTkLabel(frame, text="Selecione um arquivo Excel para iniciar:")
        label_instrucoes.pack(pady=10)

        #botao_iniciar = ctk.CTkButton(frame, text="Iniciar", command=lambda: self.trocar_tela(self.tela_formatos))
        #botao_iniciar = ctk.CTkButton(frame, text="Iniciar", command=lambda:  self.trocar_tela(self.tela_formatos)) # rode a função selecionar tabela e depois troque de tela para a tela formatos
        #botao_iniciar.pack(pady=20)
        botao_iniciar = ctk.CTkButton(frame, text="Iniciar", command=self.iniciar_processo)
        botao_iniciar.pack(pady=20)

        # Rodapé
        rodape = ctk.CTkLabel(frame, text="Desenvolvido por Júlio César.", font=ctk.CTkFont(size=11))
        rodape.pack(side='bottom', pady=10)

    def tela_gerando(self, frame):
        label_titulo = ctk.CTkLabel(frame, text="IDocumentos", font=ctk.CTkFont(size=20))
        label_titulo.pack(pady=12, padx=10)

        label_instrucoes = ctk.CTkLabel(frame, text="Gerando Arquivos...", font=ctk.CTkFont(size=20, weight="bold"))
        label_instrucoes.pack(pady=10)        

        # Rodapé
        rodape = ctk.CTkLabel(frame, text="Desenvolvido por Júlio César.", font=ctk.CTkFont(size=11))
        rodape.pack(side='bottom', pady=10)

    def iniciar_processo(self):
        # Chama a função selecionar tabela
        retorno = self.gerador.selecionar_tabela()
        if retorno == "erro":
            self.trocar_tela(self.tela_inicial)
        else:
        # Após a seleção, troca para a tela de formatos
            self.trocar_tela(self.tela_formatos)


    botao_reuniao = None
    botao_presenca = None

    def tela_formatos(self, frame):
        label_titulo = ctk.CTkLabel(frame, text="IDocumentos", font=ctk.CTkFont(size=16, weight="bold"))
        label_titulo.pack(pady=10)

        # Seção de seleção de formatos com contorno branco
        frame_formatos = ctk.CTkFrame(frame, fg_color=None, border_width=1, border_color="white")
        frame_formatos.pack(pady=10, padx=10, fill="x", expand=False)

        label_pergunta_formatos = ctk.CTkLabel(frame_formatos, text="Selecione o formato que os arquivos serão criados:")
        label_pergunta_formatos.pack(pady=10)

        self.opcao_formato = ctk.StringVar(value="DOCX")
        
        # Organizando os botões de opção lado a lado
        opcoes_frame = ctk.CTkFrame(frame_formatos, fg_color=None)
        opcoes_frame.pack(pady=5)

        radio_pdf = ctk.CTkRadioButton(opcoes_frame, text="PDF", variable=self.opcao_formato, value="PDF", command=self.verificar_selecao)
        radio_pdf.grid(row=0, column=0, padx=10)

        radio_docx = ctk.CTkRadioButton(opcoes_frame, text="DOCX", variable=self.opcao_formato, value="DOCX", command=self.verificar_selecao)
        radio_docx.grid(row=0, column=1, padx=10)

        self.label_aviso = ctk.CTkLabel(frame_formatos, text="", text_color="yellow")
        self.label_aviso.pack(pady=5)

        # Seção de seleção de ações com contorno branco
        frame_acoes = ctk.CTkFrame(frame, fg_color=None, border_width=1, border_color="white")
        frame_acoes.pack(pady=20, padx=10, fill="x", expand=False)

        label_pergunta_acoes = ctk.CTkLabel(frame_acoes, text="Selecione uma opção:")
        label_pergunta_acoes.pack(pady=10)

        #frame que armazena os button
        frame_btn = ctk.CTkFrame(frame_acoes, fg_color=None)
        frame_btn.pack(pady=20, padx=10, fill="x", expand=False)
    
        botao_irmaos = ctk.CTkButton(frame_btn, text="Procurar Irmãos", command=self.iniciar_processo1)
        botao_irmaos.grid(row=2, column=1, pady=5, padx=5)

        self.botao_reuniao = ctk.CTkButton(frame_btn, text="Arquivo Reunião", command=self.iniciar_processo2)
        self.botao_reuniao.grid(row=2, column=2, pady=5, padx=5)

        self.botao_presenca = ctk.CTkButton(frame_btn, text="Arquivo Presença",  command=lambda: self.trocar_tela(self.tela_selecao_meses))
        self.botao_presenca.grid(row=2, column=3, pady=5, padx=5)

        botao_uni = ctk.CTkButton(frame_btn, text="Arquivo uniforme", command=self.iniciar_processo4)
        botao_uni.grid(row=3, column=1, pady=5, padx=5)


            


    def tela_selecao_meses(self, frame):
        label_titulo = ctk.CTkLabel(frame, text="Gerar Arquivo Presença", font=ctk.CTkFont(size=16, weight="bold"))
        label_titulo.pack(pady=10)

        label_primeiro_mes = ctk.CTkLabel(frame, text="Digite o primeiro mês:")
        label_primeiro_mes.pack(pady=2)
        self.input_primeiro_mes = ctk.CTkEntry(frame)
        self.input_primeiro_mes.pack(pady=2)

        label_segundo_mes = ctk.CTkLabel(frame, text="Digite o segundo mês:")
        label_segundo_mes.pack(pady=2)
        self.input_segundo_mes = ctk.CTkEntry(frame)
        self.input_segundo_mes.pack(pady=2)

        botao_continuar = ctk.CTkButton(frame, text="Continuar", command=self.gerar_presenca)
        botao_continuar.pack(pady=20)

    def gerar_presenca(self):
        primeiro_mes = str(self.input_primeiro_mes.get())
        segundo_mes = str(self.input_segundo_mes.get())
        
        # Validação básica
        if not primeiro_mes or not segundo_mes:
            messagebox.showerror("Erro", "Por favor, preencha os dois meses.")
            return

        teste = self.gerador.pasta_destino()
        self.trocar_tela(self.tela_gerando)
        if teste == "erro":
            return
        if self.opcao_formato.get() == "PDF":
            self.gerador.gerarArquivoPorTurma("templates/template_Presenca.docx", None, True, primeiro_mes, segundo_mes)
        else:
            self.gerador.gerarArquivoPorTurma("templates/template_Presenca.docx", None, False, primeiro_mes, segundo_mes)
        messagebox.showinfo("Ação finalizada", f"Arquivos gerados")
        self.trocar_tela(self.tela_formatos)

    def verificar_selecao(self):
        if self.opcao_formato.get() == "PDF":
            self.label_aviso.configure(text="É necessário ter o Microsoft Word instalado para utilizar esta opção.")
            self.botao_reuniao.configure(state='disabled')
            self.botao_presenca.configure(state='disabled')
        else:
            self.label_aviso.configure(text="")
            self.botao_reuniao.configure(state='enable')
            self.botao_presenca.configure(state='enable')
    
    def verificar_selecao2(self): 
        if self.opcao_formato.get() == "PDF":
            self.botao_reuniao.configure(state='disabled')
            self.botao_presenca.configure(state='disabled')
        else:
            self.botao_reuniao.configure(state='enable')
            self.botao_presenca.configure(state='enable')

    def executar_processo(self):
        # Simula a ação com uma barra de progresso
        self.trocar_tela(self.tela_progresso)
    
    def iniciar_processo1(self):  # irmão
        messagebox.showinfo("Selecione uma Pasta", "Clique em 'OK' e escolha uma pasta onde os arquivos serão salvos.")
        teste = self.gerador.pasta_destino()
        self.trocar_tela(self.tela_gerando)
        if teste == "erro":
            messagebox.showerror("Erro", "Não foi possível selecionar a pasta. Tente novamente.")
            return
        messagebox.showinfo("Processo Iniciado", "Clique em 'OK' e aguarde enquanto os arquivos estão sendo gerados. Você receberá uma notificação quando o processo for concluído.")
        if self.opcao_formato.get() == "PDF":
            self.gerador.gerarArquivoIrmaos("templates/template_irmaos.docx", True)
        else:
            self.gerador.gerarArquivoIrmaos("templates/template_irmaos.docx", False)
        messagebox.showinfo("Ação Finalizada", "Os arquivos foram gerados com sucesso.")
        self.trocar_tela(self.tela_formatos)
        #self.trocar_tela(self.tela_progresso())

    def iniciar_processo2(self): # reunião
        messagebox.showinfo("Selecione uma Pasta", "Clique em 'OK' e escolha uma pasta onde os arquivos serão salvos.")
        teste = self.gerador.pasta_destino()
        self.trocar_tela(self.tela_gerando)
        if teste == "erro":
            messagebox.showerror("Erro", "Não foi possível selecionar a pasta. Tente novamente.")
            return
        messagebox.showinfo("Processo Iniciado", "Clique em 'OK' e aguarde enquanto os arquivos estão sendo gerados. Você receberá uma notificação quando o processo for concluído.")
        if self.opcao_formato.get() == "PDF":
            self.gerador.gerarArquivoPorTurma("templates/template_reuniao.docx", None, True, None, None)
        else:
            self.gerador.gerarArquivoPorTurma("templates/template_reuniao.docx", None, False, None, None)
        messagebox.showinfo("Ação Finalizada", "Os arquivos foram gerados com sucesso.")
        self.trocar_tela(self.tela_formatos)
        #self.trocar_tela(self.tela_progresso())
    
    def iniciar_processo3(self): # presença
        messagebox.showinfo("Selecione uma Pasta", "Clique em 'OK' e escolha uma pasta onde os arquivos serão salvos.")
        teste = self.gerador.pasta_destino()
        self.trocar_tela(self.tela_gerando)
        if teste == "erro":
            messagebox.showerror("Erro", "Não foi possível selecionar a pasta. Tente novamente.")
            return
        messagebox.showinfo("Processo Iniciado", "Clique em 'OK' e aguarde enquanto os arquivos estão sendo gerados. Você receberá uma notificação quando o processo for concluído.")
        if self.opcao_formato.get() == "PDF":
            self.gerador.gerarArquivoPorTurma("templates/template_Presenca.docx", None, True, None, None)
        else:
            self.gerador.gerarArquivoPorTurma("templates/template_Presenca.docx", None, False, None, None)
        messagebox.showinfo("Ação Finalizada", "Os arquivos foram gerados com sucesso.")
        self.trocar_tela(self.tela_formatos)
        #self.trocar_tela(self.tela_progresso())
    
    def iniciar_processo4(self):
        messagebox.showinfo("Selecione uma Pasta", "Clique em 'OK' e escolha uma pasta onde os arquivos serão salvos.")
        teste = self.gerador.pasta_destino()
        self.trocar_tela(self.tela_gerando)
        if teste == "erro":
            messagebox.showerror("Erro", "Não foi possível selecionar a pasta. Tente novamente.")
            return
        messagebox.showinfo("Processo Iniciado", "Clique em 'OK' e aguarde enquanto os arquivos estão sendo gerados. Você receberá uma notificação quando o processo for concluído.")
        if self.opcao_formato.get() == "PDF":
            self.gerador.gerarArquivoPorAluno("templates/template_Uniforme.docx", None, True)
        else:
            self.gerador.gerarArquivoPorAluno("templates/template_Uniforme.docx", None, False)
        messagebox.showinfo("Ação Finalizada", "Os arquivos foram gerados com sucesso.")
        self.trocar_tela(self.tela_formatos)
        #self.trocar_tela(self.tela_progresso())

    def tela_progresso(self, frame):
        label_progresso = ctk.CTkLabel(frame, text="Processando...", font=ctk.CTkFont(size=16, weight="bold"))
        label_progresso.pack(pady=20)

        # Barra de progresso
        self.progressbar = ctk.CTkProgressBar(frame)
        self.progressbar.pack(pady=20, padx=10)
        self.progressbar.set(0)  # Valor inicial da barra de progresso

        self.after(500, self.simular_progresso)  # Simula a ação do processo

    def simular_progresso(self):
        # Simula o avanço da barra de progresso
        progresso_atual = self.progressbar.get()
        if progresso_atual < 1:
            self.progressbar.set(progresso_atual + 0.1)
            self.after(500, self.simular_progresso)
        else:
            messagebox.showinfo("Concluído", "Processo finalizado!")
            self.trocar_tela(self.tela_inicial)  # Volta à tela inicial após o término

if __name__ == "__main__":
    #logo = ctk.CTkImage(file="templates/logo.png")
    app = App()
    tela = Tela()
    app.objeto(tela)
    tela.objeto(app)
    tela.mainloop()

